#!/usr/bin/env python3
"""
GARAMe Manager 매뉴얼 변환 스크립트 (개선 버전)
Markdown → DOCX (겉표지, 목차, 이미지 포함)

필요 라이브러리:
    pip3 install python-docx python-pptx markdown pillow
"""

import re
import os
from pathlib import Path
from datetime import datetime


def install_dependencies():
    """필요한 라이브러리 설치"""
    import subprocess
    import sys

    required = ['python-docx', 'pillow']

    print("필요한 라이브러리 확인 중...")
    for pkg in required:
        try:
            if pkg == 'python-docx':
                import docx
            elif pkg == 'pillow':
                from PIL import Image
            print(f"✓ {pkg} 설치됨")
        except ImportError:
            print(f"✗ {pkg} 미설치 - 설치 중...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", pkg])
            print(f"✓ {pkg} 설치 완료")


def parse_markdown(md_file):
    """Markdown 파일 파싱"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # 구조 분석
    sections = []
    current_section = None
    current_content = []

    for line in content.split('\n'):
        # 제목 감지
        h1_match = re.match(r'^# (.+)$', line)
        h2_match = re.match(r'^## (.+)$', line)
        h3_match = re.match(r'^### (.+)$', line)
        h4_match = re.match(r'^#### (.+)$', line)

        if h1_match:
            if current_section:
                current_section['content'] = '\n'.join(current_content)
                sections.append(current_section)
            current_section = {'level': 1, 'title': h1_match.group(1), 'content': ''}
            current_content = []
        elif h2_match:
            if current_section:
                current_section['content'] = '\n'.join(current_content)
                sections.append(current_section)
            current_section = {'level': 2, 'title': h2_match.group(1), 'content': ''}
            current_content = []
        elif h3_match:
            if current_section:
                current_section['content'] = '\n'.join(current_content)
                sections.append(current_section)
            current_section = {'level': 3, 'title': h3_match.group(1), 'content': ''}
            current_content = []
        elif h4_match:
            if current_section:
                current_section['content'] = '\n'.join(current_content)
                sections.append(current_section)
            current_section = {'level': 4, 'title': h4_match.group(1), 'content': ''}
            current_content = []
        else:
            current_content.append(line)

    # 마지막 섹션 추가
    if current_section:
        current_section['content'] = '\n'.join(current_content)
        sections.append(current_section)

    return sections


def add_cover_page(doc, logo_path=None):
    """겉표지 추가"""
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # 로고 추가 (있는 경우)
    if logo_path and os.path.exists(logo_path):
        try:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(logo_path, width=Inches(2.5))
            doc.add_paragraph()  # 여백
        except Exception as e:
            print(f"⚠️ 로고 삽입 실패: {e}")

    # 제목
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GARAMe Manager")
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = RGBColor(44, 62, 80)  # #2C3E50

    # 부제목
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("사용자 매뉴얼")
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(52, 73, 94)  # #34495E

    doc.add_paragraph()
    doc.add_paragraph()

    # 버전 정보
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run("Version 1.9.1")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(127, 140, 141)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # 설명
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("산업안전 모니터링 시스템\n다중 센서 통합 관리 플랫폼")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(127, 140, 141)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # 날짜
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date.add_run(datetime.now().strftime("%Y년 %m월"))
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(127, 140, 141)

    doc.add_paragraph()
    doc.add_paragraph()

    # 저작권
    copyright_p = doc.add_paragraph()
    copyright_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = copyright_p.add_run("Copyright © 2025 GARAMe Project")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(149, 165, 166)

    # 페이지 나누기
    doc.add_page_break()


def add_table_of_contents(doc, sections):
    """목차 추가"""
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # 목차 제목
    toc_title = doc.add_heading("목차", level=1)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.runs[0]
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph()

    # 목차 항목 (Level 1, 2만)
    for i, section in enumerate(sections):
        if section['level'] <= 2:
            title = section['title']

            # Copyright, 목차 등은 제외
            if any(skip in title for skip in ['Copyright', '목차', '©']):
                continue

            p = doc.add_paragraph(style='List Bullet' if section['level'] == 2 else 'Normal')

            # 들여쓰기
            if section['level'] == 2:
                p.paragraph_format.left_indent = Inches(0.5)

            run = p.add_run(title)
            run.font.size = Pt(12) if section['level'] == 1 else Pt(11)

            if section['level'] == 1:
                run.font.bold = True
                run.font.color.rgb = RGBColor(44, 62, 80)
            else:
                run.font.color.rgb = RGBColor(52, 73, 94)

    # 페이지 나누기
    doc.add_page_break()


def find_screenshot_for_section(section_title, captures_dir):
    """섹션에 해당하는 스크린샷 찾기"""
    # 간단한 매칭 로직 - 나중에 확장 가능
    screenshots = []

    if os.path.exists(captures_dir):
        for file in sorted(os.listdir(captures_dir)):
            if file.endswith(('.png', '.jpg', '.jpeg')):
                screenshots.append(os.path.join(captures_dir, file))

    # 섹션별 이미지 매칭 (예시)
    if "메인 화면" in section_title and screenshots:
        return screenshots[:1]  # 첫 번째 스크린샷

    return []


def markdown_to_docx_enhanced(md_file, output_file, logo_path=None, captures_dir=None):
    """Markdown → DOCX 변환 (개선 버전)"""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    print(f"\n📄 DOCX 변환 시작: {md_file}")

    doc = Document()

    # 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'  # 맑은 고딕
    font.size = Pt(11)

    # 1. 겉표지 추가
    print("  ├─ 겉표지 생성 중...")
    add_cover_page(doc, logo_path)

    # 2. 매뉴얼 파싱
    print("  ├─ 매뉴얼 파싱 중...")
    sections = parse_markdown(md_file)

    # 3. 목차 추가
    print("  ├─ 목차 생성 중...")
    add_table_of_contents(doc, sections)

    # 4. 본문 추가
    print("  ├─ 본문 생성 중...")
    for idx, section in enumerate(sections):
        level = section['level']
        title = section['title']
        content = section['content'].strip()

        # Copyright 섹션은 스킵
        if 'Copyright' in title or '©' in title:
            continue

        # 제목 추가
        if level == 1:
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = heading.runs[0]
            run.font.color.rgb = RGBColor(44, 62, 80)  # #2C3E50
            run.font.size = Pt(24)
            run.font.bold = True
        elif level == 2:
            heading = doc.add_heading(title, level=2)
            run = heading.runs[0]
            run.font.color.rgb = RGBColor(52, 73, 94)  # #34495E
            run.font.size = Pt(18)
            run.font.bold = True
        elif level == 3:
            heading = doc.add_heading(title, level=3)
            run = heading.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
        elif level == 4:
            heading = doc.add_heading(title, level=4)
            run = heading.runs[0]
            run.font.size = Pt(12)
            run.font.bold = True

        # 내용 추가
        if content:
            # 코드 블록 처리
            code_block_pattern = r'```[\w]*\n(.*?)\n```'
            parts = re.split(code_block_pattern, content, flags=re.DOTALL)

            for i, part in enumerate(parts):
                if i % 2 == 0:
                    # 일반 텍스트
                    for line in part.split('\n'):
                        line_stripped = line.strip()
                        if line_stripped:
                            # 리스트 항목 처리
                            if re.match(r'^[-*]\s+', line):
                                text = line.strip()[2:]
                                p = doc.add_paragraph(text, style='List Bullet')
                            elif re.match(r'^\d+\.\s+', line):
                                text = re.sub(r'^\d+\.\s+', '', line)
                                p = doc.add_paragraph(text, style='List Number')
                            else:
                                p = doc.add_paragraph(line_stripped)

                            # 굵은 글씨, 이탤릭 처리
                            if '**' in line_stripped or '*' in line_stripped:
                                p.clear()
                                # 굵은 글씨 처리
                                parts_bold = re.split(r'\*\*([^*]+)\*\*', line_stripped)
                                for j, part_text in enumerate(parts_bold):
                                    run = p.add_run(part_text)
                                    if j % 2 == 1:  # 굵은 글씨
                                        run.bold = True
                else:
                    # 코드 블록
                    p = doc.add_paragraph(part)
                    p_format = p.paragraph_format
                    p_format.left_indent = Inches(0.5)
                    run = p.runs[0]
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    # 배경색 (회색)
                    from docx.oxml.shared import OxmlElement
                    from docx.oxml.ns import qn
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'F5F5F5')
                    p._element.get_or_add_pPr().append(shading_elm)

        # 스크린샷 삽입 (있는 경우)
        if captures_dir:
            screenshots = find_screenshot_for_section(title, captures_dir)
            for screenshot in screenshots[:2]:  # 최대 2개
                try:
                    print(f"    ├─ 이미지 삽입: {os.path.basename(screenshot)}")
                    doc.add_paragraph()
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(screenshot, width=Inches(5.5))

                    # 이미지 캡션
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = caption.add_run(f"[그림] {title}")
                    run.font.size = Pt(9)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(127, 140, 141)
                    doc.add_paragraph()
                except Exception as e:
                    print(f"    ⚠️ 이미지 삽입 실패: {e}")

        # 섹션 간 여백 (Level 1, 2만)
        if level <= 2:
            doc.add_paragraph()

        # 주요 섹션 후 페이지 나누기 (Part 단위)
        if level == 1 and '부' in title:
            doc.add_page_break()

    # 저장
    doc.save(output_file)
    print(f"✓ DOCX 저장 완료: {output_file}")


def main():
    """메인 실행 함수"""
    import sys

    print("=" * 60)
    print("GARAMe Manager 매뉴얼 변환 스크립트 (개선 버전)")
    print("=" * 60)

    # 라이브러리 설치
    try:
        install_dependencies()
    except Exception as e:
        print(f"\n❌ 라이브러리 설치 실패: {e}")
        print("\n수동 설치 방법:")
        print("  pip3 install python-docx pillow")
        return 1

    # 경로 설정
    base_dir = Path(__file__).parent
    md_file = base_dir / "GARAMe_MANAGER_사용자매뉴얼_v1.9.1.md"

    # 로고 파일 찾기
    logo_path = None
    for logo_file in ['assets/GARAMe.png', 'assets/GARAMe1.png', 'assets/logo.png']:
        logo_candidate = base_dir / logo_file
        if logo_candidate.exists():
            logo_path = str(logo_candidate)
            print(f"✓ 로고 발견: {logo_file}")
            break

    # 스크린샷 디렉토리
    captures_dir = base_dir / "captures"
    if not captures_dir.exists():
        captures_dir = None
        print("⚠️ 스크린샷 폴더를 찾을 수 없습니다")
    else:
        screenshot_count = len(list(captures_dir.glob("*.png"))) + len(list(captures_dir.glob("*.jpg")))
        print(f"✓ 스크린샷 발견: {screenshot_count}개")

    if not md_file.exists():
        print(f"\n❌ 매뉴얼 파일을 찾을 수 없습니다: {md_file}")
        return 1

    # 출력 파일
    docx_file = base_dir / "GARAMe_MANAGER_사용자매뉴얼_v1.9.1_완성본.docx"

    print(f"\n📝 입력 파일: {md_file.name}")
    print(f"📄 출력 파일: {docx_file.name}")

    # 변환 수행
    try:
        markdown_to_docx_enhanced(
            str(md_file),
            str(docx_file),
            logo_path=logo_path,
            captures_dir=str(captures_dir) if captures_dir else None
        )

        print("\n" + "=" * 60)
        print("✅ 변환 완료!")
        print("=" * 60)
        print(f"\n생성된 파일:")
        print(f"  📄 {docx_file}")
        print(f"\n포함 내용:")
        print(f"  ✓ 겉표지 (로고 포함)" if logo_path else "  ✓ 겉표지")
        print(f"  ✓ 목차")
        print(f"  ✓ 본문 내용")
        if captures_dir:
            print(f"  ✓ 스크린샷 (자동 삽입)")

        return 0

    except Exception as e:
        import traceback
        print(f"\n❌ 변환 실패: {e}")
        print(traceback.format_exc())
        return 1


if __name__ == "__main__":
    import sys
    sys.exit(main())
