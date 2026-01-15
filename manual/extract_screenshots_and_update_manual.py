#!/usr/bin/env python3
"""
매니저 화면 자료에서 스크린샷 추출 및 매뉴얼 업데이트

1. 매니저 화면 자료.docx에서 이미지 추출
2. 추출된 이미지를 screenshots/ 폴더에 저장
3. 매뉴얼에 이미지를 적절히 배치하여 새 DOCX 생성
"""

import os
import re
from pathlib import Path
from datetime import datetime


def extract_images_from_docx(docx_path, output_dir):
    """DOCX 파일에서 이미지 추출"""
    from docx import Document
    from PIL import Image
    import io

    print(f"\n📸 이미지 추출 시작: {docx_path}")

    doc = Document(docx_path)
    os.makedirs(output_dir, exist_ok=True)

    # 텍스트와 이미지 매핑
    image_info = []

    # 단락과 이미지 정보 수집
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # 이미지가 있는 단락 확인
        for run in para.runs:
            if run._element.xpath('.//w:drawing'):
                image_info.append({
                    'para_index': i,
                    'description': text if text else f"이미지 {len(image_info) + 1}",
                    'run': run
                })

    print(f"  ├─ 발견된 이미지 설명: {len(image_info)}개")

    # 이미지 파일 추출
    extracted = []
    image_rels = {}

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_rels[rel.rId] = rel

    print(f"  ├─ 이미지 파일: {len(image_rels)}개")

    # 이미지 저장
    for idx, (rid, rel) in enumerate(image_rels.items(), 1):
        try:
            image_data = rel.target_part.blob

            # 파일 확장자 결정
            content_type = rel.target_part.content_type
            if 'png' in content_type:
                ext = 'png'
            elif 'jpeg' in content_type or 'jpg' in content_type:
                ext = 'jpg'
            else:
                ext = 'png'

            # 설명 찾기
            description = "일반"
            for info in image_info:
                # 간단한 매칭 (실제로는 더 정교한 로직 필요)
                pass

            # 파일명 생성
            filename = f"screenshot_{idx:02d}.{ext}"
            filepath = os.path.join(output_dir, filename)

            # 저장
            with open(filepath, 'wb') as f:
                f.write(image_data)

            extracted.append({
                'filename': filename,
                'filepath': filepath,
                'description': image_info[idx-1]['description'] if idx <= len(image_info) else f"이미지 {idx}"
            })

            print(f"  ├─ 추출: {filename} - {extracted[-1]['description'][:50]}")

        except Exception as e:
            print(f"  ├─ ⚠️ 이미지 {idx} 추출 실패: {e}")

    print(f"  └─ 총 {len(extracted)}개 이미지 추출 완료")

    return extracted


def create_enhanced_manual_with_screenshots(md_file, screenshots, output_file, logo_path=None):
    """스크린샷이 포함된 고급 매뉴얼 생성"""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    print(f"\n📄 매뉴얼 생성 시작")

    doc = Document()

    # 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(11)

    # 1. 겉표지
    print("  ├─ 겉표지 추가")
    add_cover_page(doc, logo_path)

    # 2. 매뉴얼 파싱
    print("  ├─ 매뉴얼 파싱")
    sections = parse_markdown(md_file)

    # 3. 목차
    print("  ├─ 목차 추가")
    add_table_of_contents(doc, sections)

    # 4. 본문 + 스크린샷
    print("  ├─ 본문 및 스크린샷 추가")

    # 스크린샷 매핑
    screenshot_mapping = {
        '첫 프로그램 실행': ['screenshot_01.png'],
        '메인 화면': ['screenshot_02.png', 'screenshot_03.png'],
        '센서 모니터링': ['screenshot_02.png', 'screenshot_04.png'],
        '거울 모드': ['screenshot_05.png'],
        '센서 통신 끊김': ['screenshot_06.png'],
        '임계값 설정': ['screenshot_07.png'],
        '안전교육': ['screenshot_08.png', 'screenshot_09.png'],
        '포스터 관리': ['screenshot_10.png'],
        '도면 관리': ['screenshot_11.png'],
        '얼굴 인식': ['screenshot_12.png'],
        '안전장구 감지': ['screenshot_13.png'],
        '접근 제어': ['screenshot_14.png'],
        '도면 보기': ['screenshot_15.png'],
        '환경설정': ['screenshot_16.png'],
    }

    for section in sections:
        level = section['level']
        title = section['title']
        content = section['content'].strip()

        # Copyright 스킵
        if 'Copyright' in title or '©' in title:
            continue

        # 제목 추가
        add_section_heading(doc, title, level)

        # 내용 추가
        if content:
            add_section_content(doc, content)

        # 스크린샷 추가
        matching_screenshots = []
        for keyword, screenshot_files in screenshot_mapping.items():
            if keyword in title:
                matching_screenshots.extend(screenshot_files)
                break

        # 매칭된 스크린샷 삽입
        for screenshot_file in matching_screenshots[:2]:  # 최대 2개
            for screenshot_info in screenshots:
                if screenshot_info['filename'] == screenshot_file:
                    try:
                        add_screenshot(doc, screenshot_info)
                        print(f"    ├─ 이미지 삽입: {screenshot_file} → {title}")
                    except Exception as e:
                        print(f"    ├─ ⚠️ 이미지 삽입 실패: {e}")
                    break

        # 섹션 간 여백
        if level <= 2:
            doc.add_paragraph()

    # 5. 스크린샷 모음 섹션 (별도)
    print("  ├─ 스크린샷 모음 섹션 추가")
    add_screenshots_appendix(doc, screenshots)

    # 저장
    doc.save(output_file)
    print(f"  └─ 저장 완료: {output_file}")


def add_cover_page(doc, logo_path):
    """겉표지"""
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if logo_path and os.path.exists(logo_path):
        try:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(logo_path, width=Inches(2.5))
            doc.add_paragraph()
        except:
            pass

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GARAMe Manager")
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = RGBColor(44, 62, 80)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("사용자 매뉴얼")
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(52, 73, 94)

    doc.add_paragraph()
    doc.add_paragraph()

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run("Version 1.9.1")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(127, 140, 141)

    doc.add_paragraph()
    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("산업안전 모니터링 시스템\n다중 센서 통합 관리 플랫폼")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(127, 140, 141)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date.add_run(datetime.now().strftime("%Y년 %m월"))
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(127, 140, 141)

    doc.add_paragraph()

    copyright_p = doc.add_paragraph()
    copyright_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = copyright_p.add_run("Copyright © 2025 GARAMe Project")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(149, 165, 166)

    doc.add_page_break()


def parse_markdown(md_file):
    """Markdown 파싱"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    sections = []
    current_section = None
    current_content = []

    for line in content.split('\n'):
        h1_match = re.match(r'^# (.+)$', line)
        h2_match = re.match(r'^## (.+)$', line)
        h3_match = re.match(r'^### (.+)$', line)
        h4_match = re.match(r'^#### (.+)$', line)

        if h1_match:
            if current_section:
                current_section['content'] = '\n'.join(current_content)
                sections.append(current_section)
            current_section = {'level': 1, 'title': h1_match.group(1), 'content': ''}
            current_content = []
        elif h2_match:
            if current_section:
                current_section['content'] = '\n'.join(current_content)
                sections.append(current_section)
            current_section = {'level': 2, 'title': h2_match.group(1), 'content': ''}
            current_content = []
        elif h3_match:
            if current_section:
                current_section['content'] = '\n'.join(current_content)
                sections.append(current_section)
            current_section = {'level': 3, 'title': h3_match.group(1), 'content': ''}
            current_content = []
        elif h4_match:
            if current_section:
                current_section['content'] = '\n'.join(current_content)
                sections.append(current_section)
            current_section = {'level': 4, 'title': h4_match.group(1), 'content': ''}
            current_content = []
        else:
            current_content.append(line)

    if current_section:
        current_section['content'] = '\n'.join(current_content)
        sections.append(current_section)

    return sections


def add_table_of_contents(doc, sections):
    """목차"""
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    toc_title = doc.add_heading("목차", level=1)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.runs[0]
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph()

    for section in sections:
        if section['level'] <= 2:
            title = section['title']
            if any(skip in title for skip in ['Copyright', '목차', '©']):
                continue

            p = doc.add_paragraph(style='List Bullet' if section['level'] == 2 else 'Normal')

            if section['level'] == 2:
                p.paragraph_format.left_indent = Inches(0.5)

            run = p.add_run(title)
            run.font.size = Pt(12) if section['level'] == 1 else Pt(11)

            if section['level'] == 1:
                run.font.bold = True
                run.font.color.rgb = RGBColor(44, 62, 80)
            else:
                run.font.color.rgb = RGBColor(52, 73, 94)

    doc.add_page_break()


def add_section_heading(doc, title, level):
    """섹션 제목 추가"""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if level == 1:
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = heading.runs[0]
        run.font.color.rgb = RGBColor(44, 62, 80)
        run.font.size = Pt(24)
        run.font.bold = True
    elif level == 2:
        heading = doc.add_heading(title, level=2)
        run = heading.runs[0]
        run.font.color.rgb = RGBColor(52, 73, 94)
        run.font.size = Pt(18)
        run.font.bold = True
    elif level == 3:
        heading = doc.add_heading(title, level=3)
        run = heading.runs[0]
        run.font.size = Pt(14)
        run.font.bold = True
    elif level == 4:
        heading = doc.add_heading(title, level=4)
        run = heading.runs[0]
        run.font.size = Pt(12)
        run.font.bold = True


def add_section_content(doc, content):
    """섹션 내용 추가"""
    from docx.shared import Pt, Inches

    code_block_pattern = r'```[\w]*\n(.*?)\n```'
    parts = re.split(code_block_pattern, content, flags=re.DOTALL)

    for i, part in enumerate(parts):
        if i % 2 == 0:
            for line in part.split('\n'):
                line_stripped = line.strip()
                if line_stripped:
                    if re.match(r'^[-*]\s+', line):
                        text = line.strip()[2:]
                        doc.add_paragraph(text, style='List Bullet')
                    elif re.match(r'^\d+\.\s+', line):
                        text = re.sub(r'^\d+\.\s+', '', line)
                        doc.add_paragraph(text, style='List Number')
                    else:
                        p = doc.add_paragraph(line_stripped)

                        # 굵은 글씨 처리
                        if '**' in line_stripped:
                            p.clear()
                            parts_bold = re.split(r'\*\*([^*]+)\*\*', line_stripped)
                            for j, part_text in enumerate(parts_bold):
                                run = p.add_run(part_text)
                                if j % 2 == 1:
                                    run.bold = True
        else:
            p = doc.add_paragraph(part)
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.runs[0]
            run.font.name = 'Consolas'
            run.font.size = Pt(10)


def add_screenshot(doc, screenshot_info):
    """스크린샷 추가"""
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    filepath = screenshot_info['filepath']
    description = screenshot_info['description']

    if os.path.exists(filepath):
        doc.add_paragraph()
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(filepath, width=Inches(5.5))

        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(f"[그림] {description}")
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(127, 140, 141)
        doc.add_paragraph()


def add_screenshots_appendix(doc, screenshots):
    """부록: 스크린샷 모음"""
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_page_break()

    heading = doc.add_heading("부록: 화면 캡처 모음", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.runs[0]
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph()

    for screenshot in screenshots:
        try:
            add_screenshot(doc, screenshot)
        except Exception as e:
            print(f"    ⚠️ 부록 이미지 추가 실패: {e}")


def main():
    """메인 실행"""
    import sys

    print("=" * 60)
    print("매니저 화면 자료 추출 및 매뉴얼 업데이트")
    print("=" * 60)

    # 경로 설정
    base_dir = Path(__file__).parent
    source_docx = Path("/Users/cyber621/Desktop/회사컴퓨터백업251104/문서/매니져 화면 자료.docx")
    screenshots_dir = base_dir / "screenshots"
    md_file = base_dir / "GARAMe_MANAGER_사용자매뉴얼_v1.9.1.md"
    output_file = base_dir / "GARAMe_MANAGER_사용자매뉴얼_v1.9.1_최종본.docx"

    # 로고
    logo_path = None
    for logo_file in ['assets/GARAMe.png', 'assets/GARAMe1.png', 'assets/logo.png']:
        logo_candidate = base_dir / logo_file
        if logo_candidate.exists():
            logo_path = str(logo_candidate)
            break

    # 1. 이미지 추출
    if not source_docx.exists():
        print(f"\n❌ 소스 파일을 찾을 수 없습니다: {source_docx}")
        return 1

    screenshots = extract_images_from_docx(str(source_docx), str(screenshots_dir))

    # 2. 매뉴얼 생성
    if not md_file.exists():
        print(f"\n❌ 매뉴얼 파일을 찾을 수 없습니다: {md_file}")
        return 1

    create_enhanced_manual_with_screenshots(
        str(md_file),
        screenshots,
        str(output_file),
        logo_path=logo_path
    )

    print("\n" + "=" * 60)
    print("✅ 완료!")
    print("=" * 60)
    print(f"\n생성된 파일:")
    print(f"  📄 {output_file}")
    print(f"\n추출된 스크린샷:")
    print(f"  📁 {screenshots_dir}/ ({len(screenshots)}개)")

    return 0


if __name__ == "__main__":
    import sys
    sys.exit(main())
